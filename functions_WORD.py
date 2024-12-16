from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from data_table import celulasTabelasOS
from functions_EXCEL import *
from docx.enum.text import WD_BREAK
from openpyxl import load_workbook
import os
import pathlib
from docx.shared import Inches

corAceitavel = "#92D050"
corAlerta = "#FFFF00"
corCritico = "#FF0000"
corNormal = "#0070C0"

arquivos = os.listdir('./')
for arquivo in arquivos:
    if arquivo.endswith(".xlsm") or arquivo.endswith(".xlsx"):
            ARQUIVO_EXCEL = arquivo
g = open(ARQUIVO_EXCEL, 'rb')
ws = load_workbook(g)

planilhas_OS = EXCEL_organizarNumerosOS(ws["Listagem"])

def WORD_deletarParagrafo(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def WORD_formatarCabecalho(celula):
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_VERTICAL.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Arial"
    cell_xml_element = celula._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#A6A6A6")
    table_cell_properties.append(shade_obj)    

def WORD_formatarCabecalhoEsquerda(celula):
    for paragraph in celula.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Arial"
    cell_xml_element = celula._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#A6A6A6")
    table_cell_properties.append(shade_obj)    

def WORD_formatarCelula(celula):
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Arial"
    cor = ''
    if celula.text == "Aceitável":
        cor = corAceitavel  
    if celula.text == "Alerta":
        cor = corAlerta  
    if celula.text == "Crítico":
        cor = corCritico
    if celula.text == "Normal":
        cor = corNormal
    cell_xml_element = celula._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), cor)
    table_cell_properties.append(shade_obj)

def WORD_formatarCelulaEsquerda(celula):
    for paragraph in celula.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"

def WORD_formatarCelulaConjunto(celula):
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Arial"

def WORD_mesclarCelulas(table, cel1:list, cel2:list, texto, format):
        table.cell(cel1[0], cel1[1]).merge(table.cell(cel2[0], cel2[1]))
        table.cell(cel1[0], cel1[1]).text = str(texto)
        if format == 1:
            WORD_formatarCelula(table.cell(cel1[0], cel1[1]))
        elif format == 2:
            WORD_formatarCabecalho(table.cell(cel1[0], cel1[1]))
        elif format == 3:
            WORD_formatarCabecalhoEsquerda(table.cell(cel1[0], cel1[1]))
        elif format == 4:
            WORD_formatarCelulaEsquerda(table.cell(cel1[0], cel1[1]))

def WORD_addTabelaParagrafo(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
    
def WORD_addParagrafoDepois(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text).add_break(break_type=WD_BREAK.PAGE)
    if text:
        new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para

def WORD_criarTabelasOS(documento):
    print("⚙ -> Adicionando tabelas de OS no WORD...")
    for paragrafo in documento.paragraphs:
        if paragrafo.text == "[tabelasOS]":
            for i in planilhas_OS:
                table = documento.add_table(rows=42, cols=12)
                p = documento.add_paragraph()
                WORD_addParagrafoDepois(paragrafo, p.text)
                WORD_addTabelaParagrafo(table, paragrafo)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            WORD_deletarParagrafo(paragrafo)
    print("✔ -> Tabelas de OS adicionadas no WORD!")

def WORD_criarTabelaListagem(planilha, documento):
    print("⚙ -> Adicionando tabela de listagem no WORD...")
    for paragrafo in documento.paragraphs:
        if paragrafo.text == "[tabela_listagem]":
            tableListagem = documento.add_table(rows=0, cols=7)
            WORD_addTabelaParagrafo(tableListagem, paragrafo)
            p = documento.add_paragraph()
            # WORD_addParagrafoDepois(paragrafo, p.text)
            tableListagem.style = 'Table Grid'
            tableListagem.alignment = WD_TABLE_ALIGNMENT.CENTER
            linhas = 0
            for row in planilha.rows:
                if EXCEL_pegarValorTabelaListagem(planilha, "A", linhas) == "":
                    break
                celulaTBLword = tableListagem.add_row().cells
                celulaTBLword[0].text = EXCEL_pegarValorTabelaListagem(planilha, "A", linhas) 
                celulaTBLword[1].text = EXCEL_pegarValorTabelaListagem(planilha, "B", linhas)
                celulaTBLword[2].text = EXCEL_pegarValorTabelaListagem(planilha, "C", linhas)
                celulaTBLword[3].text = EXCEL_pegarValorTabelaListagem(planilha, "D", linhas)
                celulaTBLword[4].text = EXCEL_pegarValorTabelaListagem(planilha, "E", linhas)
                celulaTBLword[5].text = EXCEL_pegarValorTabelaListagem(planilha, "F", linhas)
                celulaTBLword[6].text = EXCEL_pegarValorTabelaListagem(planilha, "G", linhas)
                WORD_formatarCabecalho(celulaTBLword[0])
                WORD_formatarCabecalho(celulaTBLword[1])
                if linhas == 0:
                    WORD_formatarCabecalho(celulaTBLword[2])
                    WORD_formatarCabecalho(celulaTBLword[3])
                    WORD_formatarCabecalho(celulaTBLword[4])
                    WORD_formatarCabecalho(celulaTBLword[5])
                    WORD_formatarCabecalho(celulaTBLword[6])
                if linhas != 0:
                    WORD_formatarCelulaEsquerda(celulaTBLword[2])
                    WORD_formatarCelulaConjunto(celulaTBLword[3])
                    WORD_formatarCelula(celulaTBLword[4])
                    WORD_formatarCelula(celulaTBLword[5])
                    WORD_formatarCelula(celulaTBLword[6])
                linhas += 1
            paragrafo.clear()
    print("✔ -> Tabela de listagem adicionado no arquivo WORD!")

def WORD_addValoresTabelaOS(planilha, documento, data):
    print("\n⚙ -> Adicionando valores nas tabelas de OS no WORD...")
    count = 0
    countTabelas = 2
    if len(planilhas_OS) == 0:
        return
    for table in documento.tables:
        if countTabelas != 0:
            countTabelas -= 1
            continue
        print(f"🛠 -> Adicionando valor na tabela {count+1}")
        if count == len(planilhas_OS):
            break
        for item in celulasTabelasOS:
            if item == "OS-valor":
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], f"0{count+1}" if count < 10 else count+1, celulasTabelasOS[item][3])
            elif item == "Data-valor":
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], data, celulasTabelasOS[item][3])
            elif item == "Status-valor":
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], planilhas_OS[count][1][0], celulasTabelasOS[item][3])
            elif item == "Area-valor":
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], planilhas_OS[count][1][1], celulasTabelasOS[item][3])
            elif item == "Tag-valor":
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], planilhas_OS[count][1][2], celulasTabelasOS[item][3])
            elif item == "Equipamento-valor":
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], planilhas_OS[count][1][3], celulasTabelasOS[item][3])
            else:
                WORD_mesclarCelulas(table, celulasTabelasOS[item][0], celulasTabelasOS[item][1], celulasTabelasOS[item][2], celulasTabelasOS[item][3])
        count += 1
    print("✔ -> Valores adicionados nas tabelas de OS!")

def WORD_addGraficos(paragrafo, nm):
    paragrafo.text = ''
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img = paragrafo.add_run()
    img.add_picture(rf"{str(pathlib.Path().resolve())}\chart{nm}.png", width=Inches(5))
    print("✔ -> Gráficos adicionados no arquivo WORD!")